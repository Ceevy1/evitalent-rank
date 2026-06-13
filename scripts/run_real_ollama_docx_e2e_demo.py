from __future__ import annotations

import json
import sys
from pathlib import Path
from statistics import mean
from time import perf_counter

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from evitalent.demo_samples import EXPECTED_EVENTS
from evitalent.extraction.grounding_validator import GroundingValidator
from evitalent.extraction.hybrid_extraction_pipeline import HybridExtractionPipeline
from evitalent.extraction.llm_client import LLMClient
from evitalent.extraction.llm_single_event_extractor import LLMSingleEventExtractor
from evitalent.extraction.llm_structure_extractor import LLMStructureExtractor
from evitalent.extraction.safety_validator import assert_redacted_text_safe
from evitalent.parser.docx_parser import DocxDocumentParser
from evitalent.privacy.redactor import redact_text
from evitalent.scoring.ranker import rank_candidates
from evitalent.settings import PROJECT_ROOT


MODEL_NAME = "evitalent-extractor:7b"
BASE_URL = "http://127.0.0.1:11434"
API_KEY = "ollama"


SAMPLES = {
    "hr": {
        "filename": "fictional_hr_resume_for_ollama.docx",
        "title": "虚构 HR 简历",
        "sensitive_values": ["蓝海岚", "13900001111", "lanhailan@example.test", "1992-03", "已婚", "28k", "35k"],
        "body": [
            "姓名：蓝海岚",
            "性别：女",
            "出生年月：1992-03",
            "婚姻状况：已婚",
            "当前薪资：28k",
            "期望薪资：35k",
            "手机：13900001111",
            "邮箱：lanhailan@example.test",
            "教育经历：脱敏大学 人力资源管理 本科",
            "工作经历一：2018.01-2021.12，虚构制造集团，人力资源主管，负责招聘配置与招聘流程优化。任职期间完成一线员工招聘120人，将招聘完成率提升至88%。",
            "工作经历二：2022.01-2025.12，虚构科技公司，人力资源经理，负责关键岗位招聘、组织发展与人才保留。半年内完成关键岗位招聘18人，将招聘完成率提升至91%，并推动核心岗位离职率下降15%。",
        ],
    },
    "production": {
        "filename": "fictional_production_resume_for_ollama.docx",
        "title": "虚构生产简历",
        "sensitive_values": ["周明澈", "13800002222", "zhoumingche@example.test", "1988-07", "未婚", "30k", "38k"],
        "body": [
            "姓名：周明澈",
            "性别：男",
            "出生年月：1988-07",
            "婚姻状况：未婚",
            "当前薪资：30k",
            "期望薪资：38k",
            "手机：13800002222",
            "邮箱：zhoumingche@example.test",
            "教育经历：脱敏工业大学 过程装备与控制工程 本科",
            "工作经历一：2016.03-2020.12，虚构食品工厂，生产主管，负责生产计划、质量管理与工艺优化。推动产出率提升1.2%；原料损耗下降0.6%。",
            "工作经历二：2021.01-2025.12，虚构生物制造公司，生产经理，负责自动化改造和安全合规。上线2套自动化系统；保持0质量安全事故。",
        ],
    },
    "ecommerce": {
        "filename": "fictional_ecommerce_resume_for_ollama.docx",
        "title": "虚构电商简历",
        "sensitive_values": ["秦若川", "13700003333", "qinruochuan@example.test", "1990-11", "已婚", "32k", "42k"],
        "body": [
            "姓名：秦若川",
            "性别：男",
            "出生年月：1990-11",
            "婚姻状况：已婚",
            "当前薪资：32k",
            "期望薪资：42k",
            "手机：13700003333",
            "邮箱：qinruochuan@example.test",
            "教育经历：脱敏财经大学 市场营销 本科",
            "工作经历一：2017.06-2021.12，虚构电商品牌公司，电商运营经理，负责平台运营、投放优化和内容转化。8个月GMV达到1亿元；转化率提升至12%。",
            "工作经历二：2022.01-2025.12，虚构新零售公司，电商负责人，负责多平台增长与投放效率优化。ROI达到1.5。",
        ],
    },
}


class CountingOllamaClient(LLMClient):
    def __init__(self) -> None:
        super().__init__(
            provider="local_ollama",
            base_url=BASE_URL,
            api_key=API_KEY,
            model=MODEL_NAME,
            temperature=0,
            timeout_seconds=120,
            max_retries=0,
            seed=9,
        )
        self.request_count = 0
        self.request_seconds: list[float] = []
        self.structure_request_count = 0
        self.single_event_request_count = 0
        self.current_request_kind = "unknown"

    def generate_json(self, system_prompt: str, user_prompt: str) -> dict:
        if "宏观结构" in system_prompt:
            self.current_request_kind = "structure"
        else:
            self.current_request_kind = "single_event"
        started = perf_counter()
        try:
            return super().generate_json(system_prompt, user_prompt)
        finally:
            elapsed = perf_counter() - started
            self.request_count += 1
            self.request_seconds.append(elapsed)
            if self.current_request_kind == "structure":
                self.structure_request_count += 1
            elif self.current_request_kind == "single_event":
                self.single_event_request_count += 1


def _write_docx(path: Path, title: str, body: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body:
        doc.add_paragraph(line)
    doc.save(path)


def ensure_fixture_docx_files() -> dict[str, Path]:
    out_dir = PROJECT_ROOT / "data" / "fixtures" / "source_documents"
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    for domain, spec in SAMPLES.items():
        path = out_dir / spec["filename"]
        _write_docx(path, spec["title"], spec["body"])
        paths[domain] = path
    return paths


def _redaction_passed(redacted_text: str, sensitive_values: list[str]) -> bool:
    return not any(value in redacted_text for value in sensitive_values)


def _metrics(domain: str, events) -> dict:
    expected = EXPECTED_EVENTS[domain]
    actual = [(event.event_type, event.direction, event.metric_value) for event in events]
    return {
        "event_recall": min(len(actual), len(expected)) / len(expected),
        "numeric_exact_match": sum(1 for a, e in zip(actual, expected) if a[2] == e[2]) / len(expected),
        "normalization_accuracy": sum(1 for a, e in zip(actual, expected) if a[0] == e[0] and a[1] == e[1]) / len(expected),
        "grounding_pass_rate": sum(1 for event in events if event.grounding_status == "passed") / len(expected),
    }


def main() -> None:
    client = CountingOllamaClient()
    health = client.health_check()
    if not health.ok:
        print("ollama_connected=false")
        print(f"failure_reason={health.message}")
        raise SystemExit(1)

    fixture_paths = ensure_fixture_docx_files()
    started_all = perf_counter()
    document_summaries = []
    metric_rows = []
    sensitive_leakage_count = 0

    for domain, path in fixture_paths.items():
        parsed = DocxDocumentParser().parse(path)
        redaction = redact_text(parsed.cleaned_text)
        redacted_text = redaction.redacted_text
        redaction_ok = _redaction_passed(redacted_text, SAMPLES[domain]["sensitive_values"])
        sensitive_leakage_count += 0 if redaction_ok else 1
        assert_redacted_text_safe(redacted_text)

        pipeline = HybridExtractionPipeline(
            structure_extractor=LLMStructureExtractor(client=client, use_llm=True),
            single_event_extractor=LLMSingleEventExtractor(client=client, use_llm=True),
            grounding_validator=GroundingValidator(),
        )
        doc_started = perf_counter()
        result = pipeline.extract(redacted_text, f"ollama_doc_{domain}", f"ollama_candidate_{domain}")
        inference_seconds = round(perf_counter() - doc_started, 4)
        ranking = rank_candidates([result.candidate_extraction], domain)
        item = ranking.candidates[0]
        m = _metrics(domain, result.normalized_events)
        metric_rows.append(m)
        schema_passed = result.candidate_extraction is not None
        grounding_passed = all(event.grounding_status == "passed" for event in result.normalized_events)
        safety_passed = redaction_ok and sensitive_leakage_count == 0
        summary = {
            "document_id": f"ollama_doc_{domain}",
            "domain": domain,
            "extraction_mode": "local_ollama",
            "model_name": MODEL_NAME,
            "redaction_passed": redaction_ok,
            "detected_achievement_candidate_count": len(result.achievement_candidates),
            "raw_event_count": len(result.raw_events),
            "normalized_event_count": len(result.normalized_events),
            "grounded_event_count": sum(1 for event in result.normalized_events if event.grounding_status == "passed"),
            "safety_passed": safety_passed,
            "eligible_for_scoring": result.summary["eligible_for_scoring"],
            "bcs": item.bcs,
            "eci": item.eci,
            "penalty": item.penalty,
            "rank_score": item.rank_score,
            "inference_seconds": inference_seconds,
            "schema_passed": schema_passed,
            "grounding_passed": grounding_passed,
            **m,
        }
        document_summaries.append(summary)
        print(json.dumps({k: summary[k] for k in [
            "document_id", "domain", "extraction_mode", "model_name", "redaction_passed",
            "detected_achievement_candidate_count", "raw_event_count", "normalized_event_count",
            "grounded_event_count", "safety_passed", "eligible_for_scoring", "bcs", "eci",
            "penalty", "rank_score", "inference_seconds"
        ]}, ensure_ascii=False))

    total_inference = round(perf_counter() - started_all, 4)
    aggregate = {
        "extraction_mode": "local_ollama",
        "provider": "local_ollama",
        "model_name": MODEL_NAME,
        "seed": 9,
        "ollama_connected": True,
        "actual_llm_request_count": client.request_count,
        "structure_llm_request_count": client.structure_request_count,
        "single_event_llm_request_count": client.single_event_request_count,
        "total_inference_seconds": total_inference,
        "average_llm_request_seconds": round(mean(client.request_seconds), 4) if client.request_seconds else 0.0,
        "used_mock_response": False,
        "used_cached_response": False,
        "redaction_passed": all(item["redaction_passed"] for item in document_summaries),
        "schema_passed": all(item["schema_passed"] for item in document_summaries),
        "grounding_passed": all(item["grounding_passed"] for item in document_summaries),
        "safety_passed": sensitive_leakage_count == 0,
        "eligible_for_scoring": all(item["eligible_for_scoring"] for item in document_summaries),
        "schema_valid_rate": mean(1.0 if item["schema_passed"] else 0.0 for item in document_summaries),
        "event_recall": mean(row["event_recall"] for row in metric_rows),
        "numeric_exact_match": mean(row["numeric_exact_match"] for row in metric_rows),
        "normalization_accuracy": mean(row["normalization_accuracy"] for row in metric_rows),
        "grounding_pass_rate": mean(row["grounding_pass_rate"] for row in metric_rows),
        "sensitive_leakage_count": sensitive_leakage_count,
        "documents": document_summaries,
    }
    out_dir = PROJECT_ROOT / "data" / "outputs" / "audit_reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    output_path = out_dir / "real_ollama_docx_e2e_summary.json"
    output_path.write_text(json.dumps(aggregate, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: aggregate[k] for k in [
        "ollama_connected", "actual_llm_request_count", "structure_llm_request_count",
        "single_event_llm_request_count", "total_inference_seconds", "average_llm_request_seconds",
        "schema_valid_rate", "event_recall", "numeric_exact_match", "normalization_accuracy",
        "grounding_pass_rate", "sensitive_leakage_count", "used_mock_response", "used_cached_response"
    ]}, ensure_ascii=False))
    print(f"saved_summary={output_path}")

    if client.request_count < 15:
        raise SystemExit("actual_llm_request_count below expected minimum 15")
    if aggregate["sensitive_leakage_count"] != 0:
        raise SystemExit("sensitive leakage detected")


if __name__ == "__main__":
    main()
