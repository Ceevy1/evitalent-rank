from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))


OUTPUT_DIR = ROOT / "data" / "fixtures" / "source_documents"


def _add_info_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _build_hr_resume(path: Path) -> None:
    document = Document()
    document.add_heading("个人信息", level=1)
    _add_info_table(
        document,
        [
            ("姓名", "陆晨"),
            ("性别", "女"),
            ("出生年月", "1992.05"),
            ("婚姻状况", "已婚"),
            ("家庭情况", "已婚已育"),
            ("籍贯", "虚构省云杉市"),
            ("现居住地", "虚构市星河路100号"),
            ("当前薪资", "28K/月"),
            ("期望薪资", "35K/月"),
            ("手机", "13900001234"),
            ("邮箱", "demo.hr@example.invalid"),
        ],
    )
    document.add_heading("教育经历", level=1)
    document.add_paragraph("2010.09-2014.06 虚构大学 人力资源管理 本科")
    document.add_heading("工作经历", level=1)
    document.add_paragraph("2018.01-2021.12 星桥数字服务有限公司 HRBP 主管，负责招聘配置、组织发展和培训体系搭建。")
    document.add_paragraph("2022.01-2025.06 云帆供应链管理有限公司 人力资源经理，负责关键岗位招聘、绩效激励和人才梯队建设。")
    document.add_heading("工作业绩", level=1)
    document.add_paragraph("半年完成关键岗位招聘 18 人，覆盖运营、供应链和财务岗位。")
    document.add_paragraph("通过招聘渠道优化，将招聘完成率提升至 91%。")
    document.add_heading("专业技能", level=1)
    document.add_paragraph("招聘配置、HRBP、人才盘点、培训体系、绩效激励。")
    document.save(path)


def _build_production_resume(path: Path) -> None:
    document = Document()
    document.add_heading("基本信息", level=1)
    _add_info_table(
        document,
        [
            ("姓名", "周砚"),
            ("性别", "男"),
            ("出生年月", "1989.11"),
            ("手机", "13800005678"),
            ("邮箱", "demo.production@example.invalid"),
        ],
    )
    document.add_heading("教育经历", level=1)
    document.add_paragraph("2008.09-2012.06 虚构理工学院 工业工程 本科")
    document.add_heading("工作经验", level=1)
    document.add_paragraph("2015.03-2019.08 松岭食品设备有限公司 生产主管，负责生产计划、SOP 和 7S 现场管理。")
    document.add_paragraph("2019.09-2025.06 北辰智能制造有限公司 生产经理，推动工艺优化、设备导入和自动化改造。")
    document.add_heading("项目经历", level=1)
    document.add_paragraph("自动化包装线改造项目：完成设备选型、上线验证和班组培训。")
    document.add_heading("工作业绩", level=1)
    document.add_paragraph("通过工艺参数优化，产出率提升 1.2%。")
    document.add_paragraph("通过原料领用复盘和过程控制，原料损耗下降 0.6%。")
    document.save(path)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    files = {
        "demo_hr_resume.docx": _build_hr_resume,
        "demo_production_resume.docx": _build_production_resume,
    }
    for filename, builder in files.items():
        path = OUTPUT_DIR / filename
        builder(path)
        print(f"generated={path}")


if __name__ == "__main__":
    main()
